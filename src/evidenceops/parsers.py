from __future__ import annotations

import csv
import io
import re
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from .models import ParsedChunk

SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".docx", ".txt", ".md", ".csv"}
QUESTION_STARTERS = re.compile(
    r"^(?:describe|explain|provide|detail|specify|list|identify|confirm|do|does|did|is|are|was|were|"
    r"how|what|when|where|which|who|can|could|will|would|please|whether|是否|请|如何|什么|描述|提供|说明|列出)",
    re.IGNORECASE,
)
QUESTION_LABEL = re.compile(r"^(?:questions?|prompt|requirement|item|query|问题|要求)\s*[:：]\s*", re.IGNORECASE)


class UnsupportedDocumentError(ValueError):
    pass


class DocumentParseError(ValueError):
    pass


def _clean(text: str) -> str:
    return re.sub(r"[ \t]+", " ", text.replace("\x00", " ")).strip()


def _chunk_lines(lines: Iterable[str], label: str, max_chars: int = 1_500) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    pending: list[str] = []
    pending_length = 0
    start = 1
    line_number = 0
    for line_number, raw in enumerate(lines, 1):
        line = _clean(raw)
        if not line:
            continue
        if pending and pending_length + len(line) + 1 > max_chars:
            chunks.append(ParsedChunk(text="\n".join(pending), page_or_sheet=f"{label} {start}-{line_number - 1}"))
            pending = []
            pending_length = 0
            start = line_number
        if not pending:
            start = line_number
        pending.append(line)
        pending_length += len(line) + 1
    if pending:
        chunks.append(ParsedChunk(text="\n".join(pending), page_or_sheet=f"{label} {start}-{line_number}"))
    return chunks


def _parse_pdf(data: bytes) -> list[ParsedChunk]:
    reader = PdfReader(io.BytesIO(data))
    chunks: list[ParsedChunk] = []
    for index, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        page_chunks = _chunk_lines(text.splitlines(), "Lines")
        for chunk in page_chunks:
            chunks.append(ParsedChunk(text=chunk.text, page_or_sheet=f"Page {index}"))
    return chunks


def _parse_docx(data: bytes) -> list[ParsedChunk]:
    document = Document(io.BytesIO(data))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return _chunk_lines(lines, "Paragraphs")


def _row_text(values: list[object], headers: list[str] | None) -> str:
    cells: list[str] = []
    for index, value in enumerate(values):
        if value is None or not str(value).strip():
            continue
        rendered = _clean(str(value))
        if headers and index < len(headers) and headers[index]:
            cells.append(f"{headers[index]}: {rendered}")
        else:
            cells.append(rendered)
    return " | ".join(cells)


def _parse_xlsx(data: bytes) -> list[ParsedChunk]:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    chunks: list[ParsedChunk] = []
    try:
        for worksheet in workbook.worksheets:
            rows = list(worksheet.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [_clean(str(value)) if value is not None else "" for value in rows[0]]
            use_headers = any(headers) and len(rows) > 1
            start = 1 if use_headers else 0
            for row_number, row in enumerate(rows[start:], start + 1):
                text = _row_text(list(row), headers if use_headers else None)
                if text:
                    chunks.append(ParsedChunk(text=text, page_or_sheet=f"Sheet {worksheet.title}, row {row_number}"))
    finally:
        workbook.close()
    return chunks


def _parse_csv(data: bytes) -> list[ParsedChunk]:
    text = data.decode("utf-8-sig", errors="replace")
    sample = text[:4_096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(io.StringIO(text), dialect)
    rows = list(reader)
    if not rows:
        return []
    headers = [_clean(value) for value in rows[0]]
    chunks: list[ParsedChunk] = []
    for row_number, row in enumerate(rows[1:], 2):
        rendered = _row_text(list(row), headers)
        if rendered:
            chunks.append(ParsedChunk(text=rendered, page_or_sheet=f"CSV row {row_number}"))
    return chunks


def parse_document(filename: str, data: bytes) -> list[ParsedChunk]:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentError(f"Unsupported file type: {extension or 'none'}")
    try:
        if extension == ".pdf":
            chunks = _parse_pdf(data)
        elif extension == ".docx":
            chunks = _parse_docx(data)
        elif extension == ".xlsx":
            chunks = _parse_xlsx(data)
        elif extension == ".csv":
            chunks = _parse_csv(data)
        else:
            text = data.decode("utf-8-sig", errors="replace")
            chunks = _chunk_lines(text.splitlines(), "Lines")
    except (UnsupportedDocumentError, DocumentParseError):
        raise
    except Exception as exc:
        raise DocumentParseError(f"Could not parse {filename}: {exc}") from exc
    if not chunks:
        raise DocumentParseError(f"No readable text found in {filename}")
    return chunks


def _question_candidates(text: str) -> list[str]:
    candidates: list[str] = []
    for line in text.splitlines():
        for segment in re.split(r"\s+\|\s+", line):
            cleaned = _clean(segment)
            labeled = bool(QUESTION_LABEL.match(cleaned))
            value = QUESTION_LABEL.sub("", cleaned)
            value = re.sub(r"^\s*(?:\d+[.)]|[-*•])\s*", "", value)
            if not value or len(value) < 5:
                continue
            if labeled or value.endswith(("?", "？")) or QUESTION_STARTERS.match(value):
                candidates.append(value)
    return candidates


def extract_questions(chunks: Iterable[ParsedChunk]) -> list[tuple[str, str]]:
    found: list[tuple[str, str]] = []
    seen: set[str] = set()
    for chunk in chunks:
        for question in _question_candidates(chunk.text):
            key = question.casefold()
            if key not in seen:
                found.append((question, chunk.page_or_sheet))
                seen.add(key)
    return found


def split_question_text(text: str) -> list[str]:
    synthetic = ParsedChunk(text=text, page_or_sheet="Manual input")
    extracted = [question for question, _ in extract_questions([synthetic])]
    if extracted:
        return extracted
    return [_clean(line) for line in text.splitlines() if len(_clean(line)) >= 5]
